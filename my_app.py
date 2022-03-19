from turtle import width
from docx import Document 
from docx.shared import Inches
from numpy import true_divide
import pyttsx3 as pts

def speak(text):
    pts.speak(text)



document = Document()

# Profile Picture

document.add_picture(
    'me.jpg', 
    width=Inches(1.0))

# name, phone number and email details 

name = input('Whats is your name? ')
speak('hello' + name + 'How are you today?')
speak(name + 'Whats your phone number?')
phone_number = input('Whats your phone number? ')
email = input('Whats your email? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email )

# about me

document.add_heading('About me')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# work experience

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company name ')
from_date = input ('From date ')
to_date = input ('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input ('Describe your experience at ' + company + ': ')

p.add_run(experience_details)

# more experiences

while True:
    has_more_experiences = input('Do you have more experiences? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company name ')
        from_date = input ('From date ')
        to_date = input ('To date ')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        experience_details = input ('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break 


# skills

document.add_heading('Skills')
skill = input ('Enter Skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill =  input('Enter Skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


#footer

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python"

# save documment 
document.save('cv.docx')

